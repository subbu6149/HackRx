"""
Document ingestion and processing module
Handles PDF, DOCX, and text file processing with metadata extraction
"""

import os
import sys
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
from dataclasses import dataclass

# Add parent directory to path for config import
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import settings

# Configure logging
logging.basicConfig(level=getattr(logging, settings.log_level))
logger = logging.getLogger(__name__)


@dataclass
class DocumentPage:
    """Represents a single page of a document with metadata"""
    content: str
    page_number: int
    document_source: str
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]] = None
    images: List[Dict[str, Any]] = None


@dataclass
class ProcessedDocument:
    """Represents a fully processed document"""
    document_id: str
    source_file: str
    pages: List[DocumentPage]
    total_pages: int
    metadata: Dict[str, Any]
    processing_stats: Dict[str, Any]


class DocumentIngestionSystem:
    """Handles document ingestion and initial processing"""
    
    def __init__(self):
        self.supported_formats = settings.supported_formats
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convert to bytes
        
    def validate_file(self, file_path: str) -> bool:
        """Validate if file can be processed"""
        try:
            file_path = Path(file_path)
            
            # Check if file exists
            if not file_path.exists():
                logger.error(f"File does not exist: {file_path}")
                return False
            
            # Check file extension
            if file_path.suffix.lower() not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_path.suffix}")
                return False
            
            # Check file size
            if file_path.stat().st_size > self.max_file_size:
                logger.error(f"File too large: {file_path.stat().st_size} bytes")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {str(e)}")
            return False
    
    def extract_pdf_content(self, file_path: str) -> List[DocumentPage]:
        """Extract content from PDF files using both PyPDF2 and pdfplumber"""
        pages = []
        
        try:
            # Use pdfplumber for better table extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text() or ""
                    
                    # Extract tables
                    tables = []
                    try:
                        page_tables = page.extract_tables()
                        if page_tables:
                            for i, table in enumerate(page_tables):
                                tables.append({
                                    "table_id": f"table_{page_num}_{i}",
                                    "data": table,
                                    "bbox": None  # pdfplumber doesn't provide bbox easily
                                })
                    except Exception as e:
                        logger.warning(f"Could not extract tables from page {page_num}: {str(e)}")
                    
                    # Create page metadata
                    metadata = {
                        "page_width": page.width,
                        "page_height": page.height,
                        "char_count": len(text),
                        "word_count": len(text.split()) if text else 0,
                        "has_tables": len(tables) > 0,
                        "table_count": len(tables)
                    }
                    
                    page_obj = DocumentPage(
                        content=text,
                        page_number=page_num,
                        document_source=file_path,
                        metadata=metadata,
                        tables=tables
                    )
                    
                    pages.append(page_obj)
                    
        except Exception as e:
            logger.error(f"Error extracting PDF content from {file_path}: {str(e)}")
            raise
        
        return pages
    
    def extract_docx_content(self, file_path: str) -> List[DocumentPage]:
        """Extract content from DOCX files"""
        try:
            doc = Document(file_path)
            
            # Combine all paragraphs into a single page for DOCX
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    "table_id": f"table_docx_{i}",
                    "data": table_data,
                    "bbox": None
                })
            
            metadata = {
                "char_count": len(content),
                "word_count": len(content.split()) if content else 0,
                "paragraph_count": len(doc.paragraphs),
                "has_tables": len(tables) > 0,
                "table_count": len(tables)
            }
            
            page = DocumentPage(
                content=content,
                page_number=1,
                document_source=file_path,
                metadata=metadata,
                tables=tables
            )
            
            return [page]
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {file_path}: {str(e)}")
            raise
    
    def extract_text_content(self, file_path: str) -> List[DocumentPage]:
        """Extract content from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "char_count": len(content),
                "word_count": len(content.split()) if content else 0,
                "line_count": len(content.splitlines())
            }
            
            page = DocumentPage(
                content=content,
                page_number=1,
                document_source=file_path,
                metadata=metadata,
                tables=[]
            )
            
            return [page]
            
        except Exception as e:
            logger.error(f"Error extracting text content from {file_path}: {str(e)}")
            raise
    
    def process_document(self, file_path: str) -> ProcessedDocument:
        """Main method to process a document"""
        start_time = time.time()
        
        # Validate file
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        file_path = Path(file_path)
        document_id = file_path.stem
        
        logger.info(f"Processing document: {file_path}")
        
        # Extract content based on file type
        if file_path.suffix.lower() == '.pdf':
            pages = self.extract_pdf_content(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            pages = self.extract_docx_content(str(file_path))
        elif file_path.suffix.lower() == '.txt':
            pages = self.extract_text_content(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Calculate processing stats
        processing_time = time.time() - start_time
        total_chars = sum(len(page.content) for page in pages)
        total_words = sum(page.metadata.get('word_count', 0) for page in pages)
        
        processing_stats = {
            "processing_time_seconds": processing_time,
            "total_characters": total_chars,
            "total_words": total_words,
            "pages_processed": len(pages),
            "file_size_bytes": file_path.stat().st_size
        }
        
        # Document metadata
        document_metadata = {
            "file_name": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_size": file_path.stat().st_size,
            "created_date": file_path.stat().st_ctime,
            "modified_date": file_path.stat().st_mtime
        }
        
        processed_doc = ProcessedDocument(
            document_id=document_id,
            source_file=str(file_path),
            pages=pages,
            total_pages=len(pages),
            metadata=document_metadata,
            processing_stats=processing_stats
        )
        
        logger.info(f"Successfully processed {file_path}: {len(pages)} pages, "
                   f"{total_words} words in {processing_time:.2f}s")
        
        return processed_doc


# Import time for processing stats
import time


def test_document_ingestion():
    """Test the document ingestion system"""
    ingestion_system = DocumentIngestionSystem()
    
    # Test with sample data
    sample_dir = Path(settings.data_dir)
    
    if sample_dir.exists():
        for file_path in sample_dir.glob("*.pdf"):
            try:
                processed_doc = ingestion_system.process_document(str(file_path))
                print(f"✅ Successfully processed: {file_path.name}")
                print(f"   Pages: {processed_doc.total_pages}")
                print(f"   Words: {processed_doc.processing_stats['total_words']}")
                print(f"   Processing time: {processed_doc.processing_stats['processing_time_seconds']:.2f}s")
                print()
            except Exception as e:
                print(f"❌ Error processing {file_path.name}: {str(e)}")
    else:
        print(f"Sample data directory not found: {sample_dir}")


if __name__ == "__main__":
    test_document_ingestion()
